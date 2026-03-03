"""
backend/parser.py
-----------------
Text extraction from PDF and DOCX files using pdfplumber and python-docx.
"""

import os
import io
import pdfplumber
from docx import Document


def extract_text_from_pdf(file_source) -> str:
    """
    Extract text from a PDF file.
    file_source: filepath string OR bytes-like object (from st.file_uploader)
    """
    text_parts = []
    try:
        if isinstance(file_source, (str, os.PathLike)):
            with pdfplumber.open(file_source) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        else:
            # Bytes / BytesIO from Streamlit uploader
            with pdfplumber.open(io.BytesIO(file_source.read())) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
    except Exception as e:
        print(f"[parser] PDF extraction error: {e}")
    return "\n".join(text_parts)


def extract_text_from_docx(file_source) -> str:
    """
    Extract text from a DOCX file.
    file_source: filepath string OR bytes-like object (from st.file_uploader)
    """
    text_parts = []
    try:
        if isinstance(file_source, (str, os.PathLike)):
            doc = Document(file_source)
        else:
            doc = Document(io.BytesIO(file_source.read()))

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
    except Exception as e:
        print(f"[parser] DOCX extraction error: {e}")
    return "\n".join(text_parts)


def extract_text(file_source, file_name: str = None) -> str:
    """
    Universal text extractor. Detects type from file_name or file object name.
    Returns raw text as string.
    """
    if file_name is None:
        # Try to get name from Streamlit UploadedFile
        file_name = getattr(file_source, "name", "unknown.pdf")

    ext = os.path.splitext(file_name)[-1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_source)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_source)
    else:
        # Fallback: try to read as plain text
        try:
            if isinstance(file_source, (str, os.PathLike)):
                with open(file_source, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()
            else:
                return file_source.read().decode("utf-8", errors="ignore")
        except Exception as e:
            print(f"[parser] Unknown format extraction error: {e}")
            return ""
